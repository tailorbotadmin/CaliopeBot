"""
CalíopeBot AI Orchestrator
Multi-agent editorial correction system with RAG, observability, and vector store.
"""

import io
import os
import re
import json
import uuid
import logging
import smtplib
from datetime import datetime
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.base import MIMEBase
from email import encoders
from concurrent.futures import ThreadPoolExecutor, TimeoutError as FuturesTimeoutError
from typing import List, Dict, Optional

from fastapi import FastAPI, UploadFile, File, Form, HTTPException, BackgroundTasks, Request
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel

from google import genai
from google.genai import types
from dotenv import load_dotenv
import requests

load_dotenv()

# ==========================================
# CONFIGURATION
# ==========================================
from app.config import get_settings

settings = get_settings()

# ==========================================
# LOGGING
# ==========================================
from app.services.logging_config import setup_logging, RequestTimingMiddleware

setup_logging(
    level=settings.LOG_LEVEL,
    json_format=settings.is_production,
)
logger = logging.getLogger(__name__)

# ==========================================
# FIREBASE
# ==========================================
import firebase_admin
from firebase_admin import credentials, firestore

try:
    firebase_admin.get_app()
except ValueError:
    if settings.FIREBASE_SERVICE_ACCOUNT_PATH and os.path.exists(settings.FIREBASE_SERVICE_ACCOUNT_PATH):
        cred = credentials.Certificate(settings.FIREBASE_SERVICE_ACCOUNT_PATH)
        firebase_admin.initialize_app(cred)
    else:
        firebase_admin.initialize_app()

db = firestore.client()

# ==========================================
# GEMINI CLIENT  
# Priority: VERTEX_API_KEY → GEMINI_API_KEY → Vertex ADC
# ==========================================
try:
    _vertex_key = settings.VERTEX_API_KEY or settings.GEMINI_API_KEY
    if _vertex_key and _vertex_key.startswith("AQ."):
        # Vertex AI API key — uses aiplatform.googleapis.com
        from google.genai import types as _gtypes
        logger.info(f"Initializing Gemini via Vertex AI API key (model: {settings.LLM_MODEL})")
        client = genai.Client(
            api_key=_vertex_key,
            http_options={"api_version": "v1"},
        )
        # Override the base URL to aiplatform
        client._api_client._http_options.base_url = "https://aiplatform.googleapis.com/"
    elif _vertex_key:
        # AI Studio key (generativelanguage.googleapis.com)
        logger.info(f"Initializing Gemini via AI Studio API key (model: {settings.LLM_MODEL})")
        client = genai.Client(api_key=_vertex_key)
    else:
        # Fallback: Vertex AI ADC
        logger.info("No API key — falling back to Vertex AI ADC")
        client = genai.Client(
            vertexai=True,
            project=settings.GCP_PROJECT_ID,
            location=settings.GCP_LOCATION,
        )
except Exception as e:
    logger.warning(f"Gemini client init failed: {e}")
    client = None


# ==========================================
# EMAIL CONFIG (feedback → xavi@tailorbot.tech)
# ==========================================
_FEEDBACK_RECIPIENT = "xavi@tailorbot.tech"
_GMAIL_USER         = os.getenv("GMAIL_USER", "")         # e.g. noreply@tailorbot.tech
_GMAIL_APP_PASS     = os.getenv("GMAIL_APP_PASSWORD", "") # Google Workspace App Password


# ==========================================
# VECTOR STORE (ChromaDB)
# ==========================================
from app.services.vector_store import EditorialVectorStore

try:
    vector_store = EditorialVectorStore(persist_dir=settings.CHROMA_PERSIST_DIR)
    logger.info("ChromaDB vector store initialized")
except Exception as e:
    logger.warning(f"ChromaDB init failed: {e}")
    vector_store = None

# ==========================================
# AI AGENTS
# ==========================================
from app.services.agents import VoiceAnalyzerAgent, CorrectorAgent, RevisorAgent, ArbiterAgent, CoherenceAgent

_vertex_key = settings.VERTEX_API_KEY or settings.GEMINI_API_KEY
voice_analyzer = VoiceAnalyzerAgent(client=client, model=settings.LLM_MODEL, vertex_api_key=_vertex_key)
corrector = CorrectorAgent(client=client, vector_store=vector_store, model=settings.LLM_MODEL, vertex_api_key=_vertex_key)
revisor   = RevisorAgent(client=client, vector_store=vector_store, model=settings.LLM_MODEL, vertex_api_key=_vertex_key)
arbiter   = ArbiterAgent(client=client, vector_store=vector_store, model=settings.LLM_MODEL, vertex_api_key=_vertex_key)
coherence_agent = CoherenceAgent(client=client, vector_store=vector_store, model=settings.LLM_MODEL, vertex_api_key=_vertex_key)

# ==========================================
# METRICS
# ==========================================
from app.services.metrics import (
    metrics_endpoint, CORRECTIONS_PROCESSED, SUGGESTIONS_ACCEPTED,
    SUGGESTIONS_REJECTED, ACTIVE_JOBS, VECTOR_QUERIES,
)

# ==========================================
# LANGUAGETOOL — self-hosted or public API
# ==========================================
import language_tool_python

class MockTool:
    def check(self, text):
        return []

# Detect if we have a self-hosted LT instance
_LT_HOST = os.environ.get("LANGUAGETOOL_URL", "").rstrip("/")
_LT_IS_SELFHOSTED = bool(_LT_HOST and _LT_HOST != "https://api.languagetool.org")

# Cache of initialized LanguageTool instances per language code
_lt_tools: dict = {}

def get_lt_tool(lang: str = "es-ES"):
    """Return a cached LanguageTool instance for the given language.
    Uses self-hosted server if LANGUAGETOOL_URL is set, otherwise falls back
    to the public API. Instances are cached per language code.
    """
    if lang in _lt_tools:
        return _lt_tools[lang]

    try:
        if _LT_IS_SELFHOSTED:
            logger.info(f"LanguageTool: connecting to self-hosted server at {_LT_HOST} for lang={lang}")
            # language_tool_python supports remote_server parameter
            lt = language_tool_python.LanguageTool(lang, remote_server=_LT_HOST)
        else:
            logger.info(f"LanguageTool: using public API for lang={lang}")
            lt = language_tool_python.LanguageToolPublicAPI(lang)
        _lt_tools[lang] = lt
        logger.info(f"LanguageTool initialized: {lang} ({'self-hosted' if _LT_IS_SELFHOSTED else 'public API'})")
        return lt
    except Exception as e:
        logger.warning(f"LanguageTool init failed for lang={lang}: {e} — using mock")
        mock = MockTool()
        _lt_tools[lang] = mock
        return mock

# LT tools are initialized lazily on first use (see get_lt_tool below).
# Pre-warming happens non-blocking in the lifespan startup event.

# Language configs: maps language code to normative body label
LANG_META = {
    "es":    {"lt_code": "es-ES", "normativa": "RAE",    "nombre": "Español (castellano)"},
    "es-ES": {"lt_code": "es-ES", "normativa": "RAE",    "nombre": "Español (España)"},
    "ca":    {"lt_code": "ca",    "normativa": "IEC",    "nombre": "Català"},
    "ca-ES": {"lt_code": "ca",    "normativa": "IEC",    "nombre": "Català (España)"},
    "en":    {"lt_code": "en-GB", "normativa": "Style",  "nombre": "English"},
    "en-GB": {"lt_code": "en-GB", "normativa": "Style",  "nombre": "English (UK)"},
    "en-US": {"lt_code": "en-US", "normativa": "Style",  "nombre": "English (US)"},
}

def resolve_lt_code(lang: str) -> str:
    """Normalize any language tag to the best LanguageTool code."""
    if not lang:
        return "es-ES"
    entry = LANG_META.get(lang) or LANG_META.get(lang.split("-")[0])
    return entry["lt_code"] if entry else lang


# ==========================================
# RAG BOOTSTRAP — load editorial criteria
# from Firestore into ChromaDB on startup
# ==========================================
def _bootstrap_rag_from_firestore():
    """Load all active editorial_criteria from Firestore into ChromaDB.
    
    Called once at startup so the RAG is always populated regardless of
    whether the ChromaDB persist dir survived the container restart.
    """
    if vector_store is None:
        logger.warning("[RAG bootstrap] vector_store not available, skipping")
        return
    try:
        orgs = db.collection("organizations").get()
        total_loaded = 0
        for org_doc in orgs:
            org_id = org_doc.id
            criteria_ref = (
                db.collection("organizations")
                .document(org_id)
                .collection("editorial_criteria")
            )
            criteria = criteria_ref.where("status", "==", "active").get()
            for c in criteria:
                data = c.to_dict()
                rule_text = data.get("description", "")
                rule_name = data.get("name", "")
                source    = data.get("source", "")
                if not rule_text:
                    continue
                full_text = rule_text
                if rule_name:
                    full_text = f"{rule_name}: {rule_text}"
                if source:
                    full_text += f" (Fuente: {source})"
                vector_store.add_editorial_rule(
                    org_id=org_id,
                    rule_id=f"firestore_{c.id}",
                    rule_text=full_text,
                    metadata={"category": data.get("category", ""), "source": source},
                )
                total_loaded += 1
        logger.info(f"[RAG bootstrap] Loaded {total_loaded} editorial rules from Firestore")
    except Exception as e:
        logger.error(f"[RAG bootstrap] Failed: {e}", exc_info=True)

_bootstrap_rag_from_firestore()


# ==========================================
# FASTAPI APP
# ==========================================
import asyncio
from contextlib import asynccontextmanager

async def _resume_stuck_books():
    """On every Cloud Run startup: find books stuck in 'processing' and requeue them.
    This auto-heals from redeploys that killed in-flight background tasks."""
    await asyncio.sleep(5)  # let Firebase + Gemini init settle
    try:
        orgs = list(db.collection("organizations").stream())
        for org in orgs:
            stuck = list(
                org.reference.collection("books")
                .where("status", "==", "processing")
                .stream()
            )
            for book_snap in stuck:
                book_id  = book_snap.id
                book_data = book_snap.to_dict() or {}
                org_id   = org.id
                author_id = book_data.get("authorId", "")

                chunks_ref = org.reference.collection("books").document(book_id).collection("chunks")
                pending = [c for c in chunks_ref.stream() if c.to_dict().get("status") == "pending"]

                if not pending:
                    logger.info(f"[startup] book={book_id} has no pending chunks — marking complete")
                    org.reference.collection("books").document(book_id).update(
                        {"status": "review_editor"}
                    )
                    continue

                logger.info(
                    f"[startup] Auto-resuming stuck book={book_id} org={org_id} "
                    f"({len(pending)} pending chunks)"
                )
                asyncio.create_task(
                    asyncio.to_thread(process_book_background, org_id, book_id, author_id)
                )
    except Exception as exc:
        logger.error(f"[startup] Error during stuck-book scan: {exc}", exc_info=True)


@asynccontextmanager
async def lifespan(app):
    # Startup — all tasks run as background jobs so the server binds to PORT immediately.
    asyncio.create_task(_resume_stuck_books())
    # Pre-warm LanguageTool connections in background (non-blocking)
    async def _prewarm_lt():
        await asyncio.sleep(10)  # wait for LT service to be ready
        try:
            await asyncio.to_thread(get_lt_tool, "es-ES")
            await asyncio.to_thread(get_lt_tool, "ca")
            logger.info("[startup] LanguageTool pre-warm complete")
        except Exception as e:
            logger.warning(f"[startup] LanguageTool pre-warm failed (non-fatal): {e}")
    asyncio.create_task(_prewarm_lt())
    yield
    # Shutdown — nothing special needed


app = FastAPI(
    title="CalíopeBot AI Orchestrator",
    version="2.0.0",
    description="Multi-agent editorial correction system with RAG and observability",
    lifespan=lifespan,
)

# Middleware
app.add_middleware(
    CORSMiddleware,
    allow_origins=settings.CORS_ORIGINS,
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)
app.add_middleware(RequestTimingMiddleware)

# Metrics endpoint
app.add_route("/metrics", metrics_endpoint)


# ==========================================
# MODELS
# ==========================================

class CorrectionRequest(BaseModel):
    textId: str
    text: str
    tenantId: str
    authorId: str

class SuggestionResponse(BaseModel):
    id: str
    originalText: str
    correctedText: str
    justification: str
    riskLevel: str
    sourceRule: str

class CorrectionResponse(BaseModel):
    textId: str
    suggestions: List[SuggestionResponse]

class IngestRequest(BaseModel):
    bookId: str
    organizationId: str
    fileUrl: str
    authorId: str
    language: str = "es"   # BCP 47: "es", "ca", "en", "en-GB" ...

class ExtractRulesRequest(BaseModel):
    organizationId: str
    originalFileUrl: str
    correctedFileUrl: str

class ExportRequest(BaseModel):
    originalText: str
    acceptedSuggestions: List[Dict]

class LearnRequest(BaseModel):
    tenantId: str
    authorId: str
    role: str
    originalText: str
    correctedText: str
    justification: str

class TrainStyleRequest(BaseModel):
    organizationId: str
    directory: str  # Path to directory with manuscript pairs

class CreateUserRequest(BaseModel):
    email: str
    password: str
    name: str
    role: str
    organizationId: str


# ==========================================
# FEEDBACK / SUPPORT ENDPOINT
# ==========================================

@app.post("/api/v1/feedback")
async def submit_feedback(
    raw_req: Request,
    message:      str = Form(...),
    subject:      str = Form(default="Feedback CalíopeBot"),
    sender_name:  str = Form(default=""),
    sender_email: str = Form(default=""),
    org_name:     str = Form(default=""),
    page_url:     str = Form(default=""),
    screenshots:  List[UploadFile] = File(default=[]),
):
    """Collect user feedback with optional screenshot attachments.
    Stores a record in Firestore and emails xavi@tailorbot.tech.
    Authentication is optional — public endpoint so any logged-in user can report.
    """
    # Best-effort auth to capture uid/role for the Firestore record
    uid = ""
    auth_header = raw_req.headers.get("Authorization", "")
    if auth_header.startswith("Bearer "):
        try:
            from firebase_admin import auth as _fauth
            decoded = _fauth.verify_id_token(auth_header.split("Bearer ")[1])
            uid = decoded.get("uid", "")
        except Exception:
            pass

    # ── Store in Firestore (always) ─────────────────────────────────────────
    feedback_doc = db.collection("feedback").document()
    feedback_doc.set({
        "message":     message,
        "subject":     subject,
        "senderName":  sender_name,
        "senderEmail": sender_email,
        "orgName":     org_name,
        "pageUrl":     page_url,
        "uid":         uid,
        "screenshotCount": len(screenshots),
        "timestamp":   firestore.SERVER_TIMESTAMP,
    })
    feedback_id = feedback_doc.id
    logger.info(f"Feedback stored: id={feedback_id} from={sender_email}")

    # ── Send email if SMTP credentials are configured ───────────────────────
    if _GMAIL_USER and _GMAIL_APP_PASS:
        try:
            msg = MIMEMultipart("mixed")
            msg["From"]     = f"CalíopeBot Soporte <{_GMAIL_USER}>"
            msg["To"]       = _FEEDBACK_RECIPIENT
            msg["Subject"]  = f"[CalíopeBot] {subject}"
            msg["Reply-To"] = sender_email or _GMAIL_USER

            now_str = datetime.utcnow().strftime("%Y-%m-%d %H:%M UTC")
            body = (
                f"📨 Nuevo feedback de CalíopeBot\n"
                f"═══════════════════════════════\n"
                f"De: {sender_name} <{sender_email}>\n"
                f"Organización: {org_name}\n"
                f"Página: {page_url}\n"
                f"Fecha: {now_str}\n"
                f"═══════════════════════════════\n\n"
                f"{message}\n\n"
                f"───────────────────────────────\n"
                f"Feedback ID: {feedback_id}\n"
                f"Capturas adjuntas: {len(screenshots)}\n"
            )
            msg.attach(MIMEText(body, "plain", "utf-8"))

            # Attach screenshots (read content once)
            for f_up in screenshots:
                content = await f_up.read()
                if not content:
                    continue
                part = MIMEBase("application", "octet-stream")
                part.set_payload(content)
                encoders.encode_base64(part)
                safe_name = re.sub(r'[^\w.\-]', '_', f_up.filename or "screenshot.png")
                part.add_header("Content-Disposition", f'attachment; filename="{safe_name}"')
                msg.attach(part)

            with smtplib.SMTP_SSL("smtp.gmail.com", 465, timeout=15) as srv:
                srv.login(_GMAIL_USER, _GMAIL_APP_PASS)
                srv.sendmail(_GMAIL_USER, _FEEDBACK_RECIPIENT, msg.as_string())

            logger.info(f"Feedback email sent: id={feedback_id}")
        except Exception as mail_err:
            # Non-fatal: feedback is already in Firestore
            logger.error(f"Feedback email failed (non-fatal): {mail_err}")
    else:
        logger.warning("GMAIL_USER/GMAIL_APP_PASSWORD not set — feedback stored in Firestore only")

    return {"status": "ok", "feedbackId": feedback_id}

# ==========================================
# ENDPOINTS
# ==========================================

@app.get("/api/v1/health")
def health():
    return {"status": "ok"}


@app.get("/")
def health_check():
    return {
        "status": "ok",
        "service": "CalíopeBot AI Orchestrator v2.0",
        "environment": settings.ENVIRONMENT,
        "vector_store": vector_store is not None,
        "llm_client": client is not None,
    }


@app.post("/api/v1/process-text", response_model=CorrectionResponse)
async def process_text(request: CorrectionRequest):
    """Main correction endpoint: Corrector → Revisor → Arbiter pipeline."""
    logger.info(f"Processing text {request.textId} for org={request.tenantId}")

    # LanguageTool (deterministic RAE)
    lt_matches = tool.check(request.text)
    lt_errors = [
        {
            "rule": m.ruleId,
            "message": m.message,
            "replacements": m.replacements[:3],
            "context": request.text[m.offset:m.offset + m.errorLength],
        }
        for m in lt_matches
    ]

    # No book-level voice profile available in single-text calls
    _empty_voice = {
        "resumen": "",
        "rasgos_clave": [],
        "instrucciones_agentes": "Sé conservador y respeta la voz del autor.",
        "ejemplos_representativos": [],
    }

    # Corrector → Revisor → conditional Arbiter
    corrections = corrector.run(request.text, request.tenantId, request.authorId, _empty_voice)

    for error in lt_errors:
        if error["replacements"]:
            corrections.append({
                "id": f"lt_{uuid.uuid4().hex[:8]}",
                "originalText": error["context"],
                "correctedText": error["replacements"][0],
                "justification": f"[RAE / LanguageTool] {error['message']} (Regla: {error['rule']})",
                "reglaAplicada": error["rule"],
                "riskLevel": "low",
            })

    reviews = revisor.run(request.text, corrections, request.tenantId, request.authorId, _empty_voice)
    review_map = {r["correctionId"]: r for r in reviews}

    approved, contested = [], []
    for c in corrections:
        rev = review_map.get(c.get("id", ""))
        decision = rev.get("decision", "aprobada") if rev else "aprobada"
        if decision == "aprobada":
            approved.append(c)
        elif decision == "modificada":
            c["correctedText"] = rev.get("correctedTextFinal") or c["correctedText"]
            approved.append(c)
        else:
            contested.append(c)

    arbiter_resolved = []
    if contested:
        arbiter_resolved = arbiter.run(request.text, contested, reviews, request.tenantId, request.authorId, _empty_voice)

    final_raw = approved + arbiter_resolved
    final_suggestions = []
    for s in final_raw:
        try:
            s.setdefault("id", f"s_{uuid.uuid4().hex[:8]}")
            s.setdefault("sourceRule", s.get("reglaAplicada", "AI"))
            final_suggestions.append(SuggestionResponse(**{
                k: s[k] for k in ("id", "originalText", "correctedText", "justification", "riskLevel", "sourceRule")
                if k in s
            }))
        except Exception:
            pass

    CORRECTIONS_PROCESSED.inc()
    return CorrectionResponse(textId=str(uuid.uuid4()), suggestions=final_suggestions)


async def process_book_background(org_id: str, book_id: str, author_id: str):
    """
    Multi-agent editorial pipeline:
      [0] VoiceAnalyzer  — one-time author style extraction
      [1] CorrectorAgent — specific rule-grounded corrections
      [2] RevisorAgent   — validates every correction
      [3] ArbiterAgent   — resolves disagreements (conditional)
      [+] LanguageTool   — deterministic RAE checks
    """
    ACTIVE_JOBS.inc()
    try:
        logger.info(f"Pipeline started: book={book_id}")
        chunks_ref = (
            db.collection("organizations").document(org_id)
            .collection("books").document(book_id).collection("chunks")
        )
        book_ref = (
            db.collection("organizations").document(org_id)
            .collection("books").document(book_id)
        )

        # ── Step 0: Voice Profile (extract once, reuse for all chunks) ─────────
        book_data = book_ref.get().to_dict() or {}
        voice_profile = book_data.get("voiceProfile")

        if not voice_profile:
            logger.info(f"Extracting voice profile for book={book_id}")
            sample_docs = list(chunks_ref.order_by("order").limit(20).stream())
            sample_paragraphs = [
                d.to_dict().get("text", "") for d in sample_docs
                if d.to_dict().get("text", "").strip()
            ]
            voice_profile = voice_analyzer.run(sample_paragraphs)
            book_ref.update({"voiceProfile": voice_profile})
            logger.info(f"Voice profile stored for book={book_id}: {voice_profile.get('rasgos_clave', [])}")
        else:
            logger.info(f"Reusing existing voice profile for book={book_id}")

        # ── Resolve LanguageTool for this book's language ──────────────────────
        book_snap_data = book_ref.get().to_dict() or {}
        book_lang = book_snap_data.get("language", "es")
        lt_code = resolve_lt_code(book_lang)
        lt_tool = get_lt_tool(lt_code)
        lang_meta = LANG_META.get(book_lang) or LANG_META.get(book_lang.split("-")[0]) or {"normativa": "RAE", "nombre": book_lang}
        logger.info(f"[book={book_id}] Language: {book_lang} → LT code: {lt_code} ({lang_meta['normativa']})")

        # ── Step 1: Collect pending chunks ────────────────────────────────────
        all_chunks = list(chunks_ref.order_by("order").stream())
        pending = [c for c in all_chunks if c.to_dict().get("status") == "pending"]

        if not pending:
            logger.warning(f"No pending chunks for book={book_id}")
            book_ref.update({"status": "review_editor", "processedChunks": 0})
            return

        processed_count = book_data.get("processedChunks", 0)  # resume from existing count (smart retry)

        for chunk in pending:
            data = chunk.to_dict()
            text = data["text"]

            # Per-chunk try/except: one bad chunk must NOT kill the whole book
            try:
                # ── LanguageTool (deterministic) — 15s timeout ────────────────
                try:
                    with ThreadPoolExecutor(max_workers=1) as _lt_exec:
                        _lt_future = _lt_exec.submit(lt_tool.check, text)
                        lt_matches = _lt_future.result(timeout=15)
                except (FuturesTimeoutError, Exception) as lt_err:
                    logger.warning(f"[book={book_id} chunk={chunk.id}] LanguageTool skipped: {lt_err}")
                    lt_matches = []

                lt_errors = [
                    {
                        "rule": m.ruleId,
                        "message": m.message,
                        "replacements": m.replacements[:3],
                        "context": text[m.offset:m.offset + m.errorLength],
                    }
                    for m in lt_matches
                ]

                # ── CorrectorAgent ────────────────────────────────────────────
                corrections = corrector.run(text, org_id, author_id, voice_profile)

                # Add LanguageTool errors as additional corrections
                def _lt_category(rule_id: str) -> str:
                    r = rule_id.upper()
                    if any(k in r for k in ["ACENTO", "TILDE", "DIACRIT"]): return "Tildes"
                    if any(k in r for k in ["PUNTUACION", "PUNCT", "COMA", "PUNTO"]): return "Puntuación"
                    if any(k in r for k in ["GRAM", "CONCORD", "VERB", "PREP", "DEQUE"]): return "Gramática"
                    if any(k in r for k in ["TYPO", "ESPACIO", "MAYUSC"]): return "Tipografía"
                    if any(k in r for k in ["FOREIGN", "EXTRAN", "ANGLICI"]): return "Extranjerismos"
                    return "Ortografía"

                for error in lt_errors:
                    if error["replacements"]:
                        corrections.append({
                            "id": f"lt_{uuid.uuid4().hex[:8]}",
                            "originalText": error["context"],
                            "correctedText": error["replacements"][0],
                            "justification": f"LanguageTool RAE — {error['message']} (Regla: {error['rule']})",
                            "reglaAplicada": error["rule"],
                            "riskLevel": "low",
                            "category": _lt_category(error["rule"]),
                        })


                for c in corrections:
                    c.setdefault("status", "pending")

                final_suggestions = []

                if corrections:
                    # ── RevisorAgent ──────────────────────────────────────────
                    reviews = revisor.run(text, corrections, org_id, author_id, voice_profile)
                    review_map = {r["correctionId"]: r for r in reviews}

                    approved, contested = [], []
                    for c in corrections:
                        rev = review_map.get(c["id"])
                        if not rev:
                            # No review — include as-is
                            approved.append(c)
                            continue
                        decision = rev.get("decision", "aprobada")
                        if decision == "aprobada":
                            approved.append(c)
                        elif decision == "modificada":
                            c["correctedText"] = rev.get("correctedTextFinal") or c["correctedText"]
                            c["justification"] += f" [Revisor: {rev.get('razon', '')}]"
                            approved.append(c)
                        else:  # rechazada
                            contested.append(c)

                    # ── ArbiterAgent (only if there are contested corrections) ─
                    if contested:
                        logger.info(f"[book={book_id} chunk={chunk.id}] Arbiter resolving {len(contested)} contested")
                        arbiter_suggestions = arbiter.run(
                            text, contested, reviews, org_id, author_id, voice_profile
                        )
                        for s in arbiter_suggestions:
                            s.setdefault("status", "pending")
                        final_suggestions = approved + arbiter_suggestions
                    else:
                        final_suggestions = approved
                # If no corrections from either agent, final_suggestions stays []

                # ── Sanity filter: remove corrections whose originalText isn't in the text ─
                # Prevents LLM hallucinations (e.g. "europa" when text has "Europa")
                valid_suggestions = [
                    s for s in final_suggestions
                    if s.get("originalText", "") and s["originalText"] in text
                    and s.get("originalText") != s.get("correctedText")  # skip no-op corrections
                ]
                if len(valid_suggestions) < len(final_suggestions):
                    dropped = len(final_suggestions) - len(valid_suggestions)
                    logger.info(f"[book={book_id} chunk={chunk.id}] Dropped {dropped} hallucinated/no-op corrections")
                final_suggestions = valid_suggestions

                # ── Commit this chunk immediately (progressive loading) ────────
                chunks_ref.document(chunk.id).update({
                    "status": "processed",
                    "suggestions": final_suggestions,
                })
                processed_count += 1
                book_ref.update({"processedChunks": processed_count})
                logger.info(f"book={book_id}: chunk {processed_count}/{len(pending)} done ({len(final_suggestions)} suggestions)")

            except Exception as chunk_err:
                # ONE chunk failure must NOT kill the whole book pipeline
                logger.error(f"[book={book_id} chunk={chunk.id}] Chunk failed, skipping: {chunk_err}", exc_info=True)
                chunks_ref.document(chunk.id).update({
                    "status": "processed",
                    "suggestions": [],
                    "_chunkError": str(chunk_err)[:300],
                })
                processed_count += 1
                book_ref.update({"processedChunks": processed_count})

        book_ref.update({"status": "review_editor", "processedChunks": processed_count})
        logger.info(f"Pipeline complete: book={book_id}, chunks={processed_count}")

        # ══ Step Final: Book-level coherence & editorial quality analysis ═════════════
        try:
            logger.info(f"[book={book_id}] Starting book-level coherence analysis...")
            all_texts_for_coherence = [
                c.to_dict().get("text", "") for c in all_chunks
                if c.to_dict().get("text", "").strip()
            ]
            full_text_for_coherence = "\n\n".join(all_texts_for_coherence)

            coherence_corrections = coherence_agent.run(
                full_text_for_coherence, org_id, voice_profile
            )

            if coherence_corrections:
                for cc in coherence_corrections:
                    original = cc.get("originalText", "").strip()
                    if not original:
                        continue
                    # Find the chunk containing this fragment (exact match)
                    target_chunk_id = None
                    for chunk_doc in all_chunks:
                        if original in chunk_doc.to_dict().get("text", ""):
                            target_chunk_id = chunk_doc.id
                            break
                    # If not found (cross-chunk issue), attach to first content chunk
                    if target_chunk_id is None and all_chunks:
                        target_chunk_id = all_chunks[0].id

                    if target_chunk_id:
                        current_data = chunks_ref.document(target_chunk_id).get().to_dict() or {}
                        current_suggs = current_data.get("suggestions", [])
                        current_suggs.append(cc)
                        chunks_ref.document(target_chunk_id).update({"suggestions": current_suggs})

                logger.info(f"[book={book_id}] Coherence: {len(coherence_corrections)} issues attached")

            # Store expanded editorial profile in the book document
            editorial_analysis = {
                "tipo_texto":          voice_profile.get("tipo_texto", ""),
                "registro":            voice_profile.get("registro", ""),
                "audiencia_objetivo":  voice_profile.get("audiencia_objetivo", ""),
                "variedad_linguistica": voice_profile.get("variedad_linguistica", ""),
                "decisiones_autorales": voice_profile.get("decisiones_autorales", []),
                "riesgos_editoriales":  voice_profile.get("riesgos_editoriales", []),
                "rasgos_clave":         voice_profile.get("rasgos_clave", []),
                "coherence_issues":     len(coherence_corrections),
            }
            book_ref.update({"editorial_analysis": editorial_analysis})
            logger.info(f"[book={book_id}] Editorial analysis stored: {editorial_analysis}")

        except Exception as coh_err:
            logger.error(f"[book={book_id}] Coherence analysis failed (non-fatal): {coh_err}", exc_info=True)


    except Exception as e:
        logger.error(f"Background task error for book={book_id}: {e}", exc_info=True)
        # Mark book as error so the user can see it and retry
        try:
            book_ref = (
                db.collection("organizations").document(org_id)
                .collection("books").document(book_id)
            )
            book_ref.update({
                "status": "error",
                "errorMessage": str(e)[:500],  # cap at 500 chars
            })
        except Exception as inner_e:
            logger.error(f"Could not update error status for book={book_id}: {inner_e}")
    finally:
        ACTIVE_JOBS.dec()


class RetryBookRequest(BaseModel):
    organizationId: str
    bookId: str
    authorId: str = ""


@app.post("/api/v1/retry-book")
async def retry_book(request: RetryBookRequest, background_tasks: BackgroundTasks, raw_req: Request):
    """Re-trigger analysis for a stuck book (all pending chunks, no re-ingestion)."""
    auth_header = raw_req.headers.get("Authorization", "")
    if auth_header.startswith("Bearer "):
        try:
            from firebase_admin import auth as fauth
            decoded = fauth.verify_id_token(auth_header.split("Bearer ")[1])
            role = decoded.get("role", "")
            if role not in ["SuperAdmin", "Responsable_Editorial", "Editor"]:
                raise HTTPException(status_code=403, detail="Forbidden")
        except HTTPException:
            raise
        except Exception:
            raise HTTPException(status_code=401, detail="Invalid token")

    org_id    = request.organizationId
    book_id   = request.bookId
    author_id = request.authorId

    book_ref = (
        db.collection("organizations").document(org_id)
        .collection("books").document(book_id)
    )
    book_snap = book_ref.get()
    if not book_snap.exists:
        raise HTTPException(status_code=404, detail="Book not found")

    chunks_ref = book_ref.collection("chunks")
    all_chunks = list(chunks_ref.stream())
    total = len(all_chunks)

    # ── Case 1: No chunks at all ─ ingestion never completed ─────────────
    if total == 0:
        raise HTTPException(
            status_code=409,
            detail="no_chunks:El manuscrito no tiene segmentos. Usa 'Reintentar' desde la lista de manuscritos para volver a subirlo."
        )

    # ── Smart retry: ONLY reset chunks still in 'pending' status ──────────
    # Chunks with status="processed" are kept regardless of whether they have
    # suggestions (empty suggestions = legitimately no errors found).
    # This avoids re-processing the entire book when Cloud Run redeploys mid-analysis.
    to_reset   = [c for c in all_chunks if c.to_dict().get("status") == "pending"]
    done_count = total - len(to_reset)

    batch_size = 450
    for i in range(0, len(to_reset), batch_size):
        batch = db.batch()
        for c in to_reset[i:i + batch_size]:
            batch.update(chunks_ref.document(c.id), {"status": "pending", "suggestions": []})
        batch.commit()


    # Reset book state — keep voiceProfile if partially done (avoid regenerating it)
    book_snap_dict = book_snap.to_dict() or {}
    update_data: dict = {
        "status": "processing",
        "processedChunks": done_count,  # resume counter from preserved chunks
        "totalChunks": total,
    }
    if done_count == 0:
        update_data["voiceProfile"] = None  # force regeneration only when starting fresh
    book_ref.update(update_data)

    background_tasks.add_task(process_book_background, org_id, book_id, author_id)
    pending_count = len(to_reset)
    logger.info(
        f"Smart retry: book={book_id}, {done_count} chunks preserved, "
        f"{pending_count} chunks reset to pending"
    )

    return {
        "status": "retrying",
        "bookId": book_id,
        "totalChunks": total,
        "pendingChunks": pending_count,
        "preservedChunks": done_count,
    }



@app.post("/api/v1/ingest-book")
async def ingest_book(request: IngestRequest, background_tasks: BackgroundTasks):
    """Parse DOCX manuscript and create chunks in Firestore."""
    try:
        from docx import Document

        response = requests.get(request.fileUrl, timeout=60)
        if response.status_code != 200:
            raise HTTPException(status_code=400, detail="Failed to download file")

        doc = Document(io.BytesIO(response.content))
        chunks_ref = (
            db.collection("organizations").document(request.organizationId)
            .collection("books").document(request.bookId).collection("chunks")
        )

        # ── Helper: detect explicit page break in a paragraph's XML ─────────
        def _para_has_page_break(para) -> bool:
            """True if the paragraph contains an explicit page break (w:br type="page")."""
            xml = para._element.xml
            return 'w:type="page"' in xml or "w:type='page'" in xml

        # ── Helper: clean numeric artifacts fused with words ─────────────────
        # Word sometimes embeds footnote/endnote markers or page numbers
        # directly in the text content, producing tokens like "Detrás75" or "también2".
        # We remove standalone digits that are fused to the START or END of a word.
        _ARTIFACT_RE = re.compile(
            r'(\b[A-Za-záéíóúüñÁÉÍÓÚÜÑ]{2,})\d+\b'   # word followed by digits: "Detrás75"
            r'|\b\d+([A-Za-záéíóúüñÁÉÍÓÚÜÑ]{2,}\b)'  # digits followed by word: "75Detrás"
        )
        def _clean_text(raw: str) -> str:
            # Remove fused numeric artifacts but preserve real numeric content
            cleaned = _ARTIFACT_RE.sub(lambda m: m.group(1) or m.group(2), raw)
            # Collapse multiple spaces
            cleaned = re.sub(r'  +', ' ', cleaned)
            return cleaned.strip()

        # ── Build chunks with page tracking ──────────────────────────────────
        all_chunks = []
        chunk_index = 0
        current_page = 1

        for para in doc.paragraphs:
            # Page break BEFORE processing this paragraph bumps the page counter
            if _para_has_page_break(para):
                current_page += 1

            raw_text = para.text.strip()
            if not raw_text:
                continue

            text = _clean_text(raw_text)
            if not text:
                continue

            all_chunks.append({
                "id": f"chunk_{str(chunk_index).zfill(4)}",
                "text": text,
                "style": para.style.name if para.style else "Normal",
                "status": "pending",
                "order": chunk_index,
                "page": current_page,
            })
            chunk_index += 1

        # Commit in batches of 450 (safe margin under 500 limit)
        for i in range(0, len(all_chunks), 450):
            batch = db.batch()
            for chunk_data in all_chunks[i:i + 450]:
                doc_ref = chunks_ref.document(chunk_data["id"])
                batch.set(doc_ref, chunk_data)
            batch.commit()

        # Update book status
        book_ref = (
            db.collection("organizations").document(request.organizationId)
            .collection("books").document(request.bookId)
        )
        book_ref.update({
            "status": "processing",
            "totalChunks": len(all_chunks),
            "processedChunks": 0,
            "totalPages": current_page,
            "language": request.language or "es",
        })


        background_tasks.add_task(
            process_book_background, request.organizationId, request.bookId, request.authorId
        )

        logger.info(f"Ingested book {request.bookId}: {len(all_chunks)} chunks across {current_page} pages")
        return {"status": "success", "total_chunks": len(all_chunks), "total_pages": current_page}

    except Exception as e:
        logger.error(f"Ingest error: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=str(e))



@app.post("/api/v1/extract-rules")
async def extract_rules(request: ExtractRulesRequest):
    """Extract editorial rules by comparing original vs corrected documents."""
    try:
        from docx import Document

        resp_orig = requests.get(request.originalFileUrl, timeout=60)
        resp_corr = requests.get(request.correctedFileUrl, timeout=60)

        if resp_orig.status_code != 200 or resp_corr.status_code != 200:
            raise HTTPException(status_code=400, detail="Failed to download files")

        doc_orig = Document(io.BytesIO(resp_orig.content))
        doc_corr = Document(io.BytesIO(resp_corr.content))

        text_orig = "\n".join([p.text for p in doc_orig.paragraphs[:100] if p.text.strip()])
        text_corr = "\n".join([p.text for p in doc_corr.paragraphs[:100] if p.text.strip()])

        if not client:
            rules = [{"id": f"p{uuid.uuid4().hex[:8]}", "rule": "Mock rule", "description": "No LLM.", "status": "pending"}]
        else:
            schema = {
                "type": "ARRAY",
                "items": {
                    "type": "OBJECT",
                    "properties": {
                        "rule": {"type": "STRING"},
                        "description": {"type": "STRING"},
                    },
                    "required": ["rule", "description"],
                },
            }
            prompt = f"""Eres un Analista Editorial experto en lengua española.
Compara estos dos textos y DEDUCE las reglas editoriales sistemáticas que aplicó el corrector.
Busca patrones consistentes. No devuelvas correcciones puntuales, sino REGLAS GENERALES.

IMPORTANTE: Responde SIEMPRE en español. Los campos 'rule' y 'description' deben estar
escritos en español, independientemente del idioma del texto analizado.

=== TEXTO ORIGINAL ===
{text_orig}

=== TEXTO CORREGIDO ===
{text_corr}"""

            response = client.models.generate_content(
                model=settings.LLM_MODEL,
                contents=prompt,
                config=types.GenerateContentConfig(
                    response_mime_type="application/json",
                    response_schema=schema,
                    temperature=0.2,
                ),
            )
            deduced = json.loads(response.text)
            rules = []
            for r in deduced:
                rules.append({
                    "id": f"p{uuid.uuid4().hex[:8]}",
                    "rule": r["rule"],
                    "description": r["description"],
                    "status": "pending",
                    "createdAt": firestore.SERVER_TIMESTAMP,
                })

        # Store in Firestore + Vector DB
        batch = db.batch()
        org_ref = db.collection("organizations").document(request.organizationId)
        for rule in rules:
            rule_ref = org_ref.collection("pendingRules").document(rule["id"])
            batch.set(rule_ref, rule)
            if vector_store:
                vector_store.add_editorial_rule(
                    request.organizationId, rule["id"],
                    f"{rule['rule']}: {rule['description']}",
                )
        batch.commit()

        logger.info(f"Extracted {len(rules)} rules for org={request.organizationId}")
        return {"status": "success", "extractedRules": rules}

    except Exception as e:
        logger.error(f"Extract rules error: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/api/v1/export-docx")
async def export_docx(request: ExportRequest):
    """Export corrected document as DOCX."""
    from docx import Document

    doc = Document()
    doc.add_paragraph("CalíopeBot - Documento Corregido")

    final_text = request.originalText
    for sug in request.acceptedSuggestions:
        final_text = final_text.replace(
            sug.get("originalText", ""), sug.get("correctedText", "")
        )
    doc.add_paragraph(final_text)

    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)

    return {"message": "Export successful", "size": len(file_stream.getvalue())}


@app.post("/api/v1/learn-correction")
async def learn_correction(request: LearnRequest):
    """Self-learning: inject accepted correction into vector store for RAG."""
    if vector_store:
        vector_store.learn_from_correction(
            org_id=request.tenantId,
            author_id=request.authorId,
            original=request.originalText,
            corrected=request.correctedText,
            justification=request.justification,
        )
        SUGGESTIONS_ACCEPTED.inc()
        VECTOR_QUERIES.labels(collection_type="learn").inc()
        logger.info(f"Learned correction for org={request.tenantId}, author={request.authorId}")
    else:
        logger.warning("Vector store not available, correction not persisted")

    return {
        "status": "success",
        "message": f"Correction learned for author {request.authorId} by {request.role}",
    }


@app.post("/api/v1/train-style")
async def train_style(request: TrainStyleRequest, background_tasks: BackgroundTasks):
    """Batch process manuscript pairs to extract editorial style rules."""
    from app.services.style_trainer import StyleTrainer

    trainer = StyleTrainer(client=client, vector_store=vector_store)

    async def _run():
        total = trainer.batch_process_manuscripts(
            request.directory, request.organizationId, db
        )
        logger.info(f"Style training complete: {total} rules extracted")

    background_tasks.add_task(_run)
    return {"status": "processing", "message": "Style training started in background"}


class UpdateRoleRequest(BaseModel):
    targetUid: str
    role: str

@app.post("/api/v1/users/update-role")
async def update_user_role(request: UpdateRoleRequest, raw_req: Request):
    """Update a user's role in Firebase Auth custom claims and Firestore."""
    auth_header = raw_req.headers.get("Authorization")
    if not auth_header or not auth_header.startswith("Bearer "):
        raise HTTPException(status_code=401, detail="Missing or invalid token")

    token = auth_header.split("Bearer ")[1]
    try:
        from firebase_admin import auth as firebase_auth
        decoded_token = firebase_auth.verify_id_token(token)

        caller_role = decoded_token.get("role", "")
        if caller_role not in ["Admin", "SuperAdmin", "Responsable_Editorial"]:
            raise HTTPException(status_code=403, detail="Forbidden, insufficient permissions")

        # Responsable_Editorial cannot promote to SuperAdmin
        if caller_role in ["Admin", "Responsable_Editorial"] and request.role in ["Admin", "SuperAdmin"]:
            raise HTTPException(status_code=403, detail="No tienes permisos para asignar ese rol")

        # Get target user's current org to prevent cross-org attacks
        target_user = db.collection("users").document(request.targetUid).get()
        if not target_user.exists:
            raise HTTPException(status_code=404, detail="User not found")

        target_org = target_user.to_dict().get("organizationId")
        caller_org = decoded_token.get("organizationId")
        if caller_role != "SuperAdmin" and target_org != caller_org:
            raise HTTPException(status_code=403, detail="Cannot edit users outside your organization")

        # Update custom claims (immediate token propagation)
        current_claims = firebase_auth.get_user(request.targetUid).custom_claims or {}
        current_claims["role"] = request.role
        firebase_auth.set_custom_user_claims(request.targetUid, current_claims)

        # Update Firestore profile
        db.collection("users").document(request.targetUid).update({"role": request.role})

        logger.info(f"Role updated: uid={request.targetUid} -> role={request.role} by {decoded_token['uid']}")
        return {"status": "success", "uid": request.targetUid, "role": request.role}

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error updating role: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/api/v1/users/create")
async def create_user_endpoint(request: CreateUserRequest, raw_req: Request):
    """Admin endpoint to create a new user and set custom claims securely."""
    auth_header = raw_req.headers.get("Authorization")
    if not auth_header or not auth_header.startswith("Bearer "):
        raise HTTPException(status_code=401, detail="Missing or invalid token")
    
    token = auth_header.split("Bearer ")[1]
    try:
        from firebase_admin import auth as firebase_auth
        decoded_token = firebase_auth.verify_id_token(token)
        
        # Only Responsable_Editorial, and SuperAdmins can create users
        if decoded_token.get("role") not in ["Admin", "SuperAdmin", "Responsable_Editorial"]:
            raise HTTPException(status_code=403, detail="Forbidden, insufficient permissions")
            
        # Security: You can't create users outside your assigned org unless you are SuperAdmin
        if decoded_token.get("role") != "SuperAdmin" and decoded_token.get("organizationId") != request.organizationId:
            raise HTTPException(status_code=403, detail="Forbidden, cannot create users outside your organization")

        user_record = firebase_auth.create_user(
            email=request.email,
            password=request.password,
            display_name=request.name
        )
        
        firebase_auth.set_custom_user_claims(user_record.uid, {
            "role": request.role,
            "organizationId": request.organizationId
        })
        
        db.collection("users").document(user_record.uid).set({
            "email": request.email,
            "name": request.name,
            "role": request.role,
            "organizationId": request.organizationId,
            "createdAt": firestore.SERVER_TIMESTAMP
        })
        
        logger.info(f"User {request.email} created successfully with role {request.role}")
        return {"status": "success", "uid": user_record.uid}
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error creating user: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=str(e))

class DeleteUserRequest(BaseModel):
    targetUid: str


@app.delete("/api/v1/users/delete")
async def delete_user(request: DeleteUserRequest, raw_req: Request):
    """Admin endpoint to delete a user from Firebase Auth and Firestore."""
    auth_header = raw_req.headers.get("Authorization")
    if not auth_header or not auth_header.startswith("Bearer "):
        raise HTTPException(status_code=401, detail="Missing or invalid token")

    token = auth_header.split("Bearer ")[1]
    try:
        from firebase_admin import auth as firebase_auth

        decoded_token = firebase_auth.verify_id_token(token)
        caller_role = decoded_token.get("role", "")

        if caller_role not in ["Admin", "SuperAdmin", "Responsable_Editorial"]:
            raise HTTPException(status_code=403, detail="Forbidden, insufficient permissions")

        # Fetch target user to enforce cross-org and self-delete protections
        target_doc = db.collection("users").document(request.targetUid).get()
        if not target_doc.exists:
            raise HTTPException(status_code=404, detail="User not found")

        target_data = target_doc.to_dict()
        target_role = target_data.get("role", "")
        target_org = target_data.get("organizationId")
        caller_org = decoded_token.get("organizationId")

        # Prevent cross-org deletion (except SuperAdmin)
        if caller_role != "SuperAdmin" and target_org != caller_org:
            raise HTTPException(status_code=403, detail="Cannot delete users outside your organization")

        # Responsable_Editorial cannot delete SuperAdmins or other Responsables
        if caller_role in ["Admin", "Responsable_Editorial"] and target_role in ["Admin", "SuperAdmin", "Responsable_Editorial"]:
            raise HTTPException(status_code=403, detail="No tienes permisos para eliminar ese usuario")

        # Prevent self-deletion
        if request.targetUid == decoded_token.get("uid"):
            raise HTTPException(status_code=400, detail="Cannot delete your own account")

        # Delete from Firebase Auth
        firebase_auth.delete_user(request.targetUid)

        # Delete Firestore profile
        db.collection("users").document(request.targetUid).delete()

        logger.info(
            f"User deleted: uid={request.targetUid} (role={target_role}) "
            f"by {decoded_token['uid']} (role={caller_role})"
        )
        return {"status": "success", "uid": request.targetUid}

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error deleting user: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=str(e))


class SeedRAERulesRequest(BaseModel):
    organizationId: str


@app.post("/api/v1/seed-rae-rules")
async def seed_rae_rules_endpoint(request: SeedRAERulesRequest, raw_req: Request):
    """SuperAdmin endpoint: populate an org with the canonical RAE rules corpus."""
    auth_header = raw_req.headers.get("Authorization")
    if not auth_header or not auth_header.startswith("Bearer "):
        raise HTTPException(status_code=401, detail="Missing or invalid token")

    token = auth_header.split("Bearer ")[1]
    try:
        from firebase_admin import auth as firebase_auth

        decoded_token = firebase_auth.verify_id_token(token)
        if decoded_token.get("role") != "SuperAdmin":
            raise HTTPException(status_code=403, detail="Only SuperAdmins can seed RAE rules")

        # Import the RAE corpus from the pure data module (no Firebase side-effects)
        from rae_rules_corpus import RAE_RULES  # noqa: PLC0415

        org_ref = db.collection("organizations").document(request.organizationId)

        # Check which rules already exist to avoid duplicates
        existing_snap = org_ref.collection("rules").stream()
        existing_ids = {doc.id for doc in existing_snap}

        batch = db.batch()
        inserted = 0
        skipped = 0

        for i, rule in enumerate(RAE_RULES):
            rule_id = f"rae_{i:03d}"
            if rule_id in existing_ids:
                skipped += 1
                continue

            rule_doc = {
                **rule,
                "status": "active",
                "source": rule.get("source", "RAE"),
                "createdAt": firestore.SERVER_TIMESTAMP,
            }
            batch.set(org_ref.collection("rules").document(rule_id), rule_doc)

            # Index in ChromaDB if available
            if vector_store:
                try:
                    vector_store.add_editorial_rule(
                        request.organizationId,
                        rule_id,
                        f"{rule['name']}: {rule['description']}",
                    )
                except Exception as ve:
                    logger.warning(f"ChromaDB indexing failed for {rule_id}: {ve}")

            inserted += 1

        batch.commit()

        logger.info(
            f"seed-rae-rules: org={request.organizationId} "
            f"inserted={inserted} skipped={skipped} "
            f"by uid={decoded_token['uid']}"
        )
        return {
            "status": "success",
            "organizationId": request.organizationId,
            "inserted": inserted,
            "skipped": skipped,
            "total": len(RAE_RULES),
        }

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error seeding RAE rules: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=str(e))


# ── Language detection helpers ───────────────────────────────────────────────

_SPANISH_PATTERN = re.compile(
    r"[ñáéíóúüÁÉÍÓÚÜ]|"
    r"\b(de|el|la|los|las|en|que|con|por|para|como|del|una|sin|este|esta|se|su|sus|al|"
    r"pero|más|no|es|son|fue|han|hay|un|su|ya|lo|le|si|yo|mi|te|me)\b",
    re.IGNORECASE,
)
_ENGLISH_PATTERN = re.compile(
    r"\b(the|and|or|with|for|from|text|using|when|should|must|will|are|is|be|that|this|which|"
    r"have|has|been|their|can|may|format|list|header|italic|bold|margin|spacing|font|style|"
    r"rule|item|use|used|applied|correction|spelling|capitalization|punctuation|alignment|"
    r"inclusion|placement|numbering|formatting|usage|dialogue|paragraph|footnote|quotation|"
    r"consistency|choice|structure|flow|marks|page|title|chapter|content)\b",
    re.IGNORECASE,
)


def _is_english_text(text: str) -> bool:
    if not text:
        return False
    return bool(_ENGLISH_PATTERN.search(text)) and not bool(_SPANISH_PATTERN.search(text))


class TranslateRulesRequest(BaseModel):
    organizationId: str


@app.post("/api/v1/translate-rules")
async def translate_rules_endpoint(request: TranslateRulesRequest, raw_req: Request):
    """SuperAdmin endpoint: detect English rules in an org and translate them to Spanish.
    Uses the Gemini client already configured in this worker — no extra credentials needed.
    """
    auth_header = raw_req.headers.get("Authorization")
    if not auth_header or not auth_header.startswith("Bearer "):
        raise HTTPException(status_code=401, detail="Missing or invalid token")

    token = auth_header.split("Bearer ")[1]
    try:
        from firebase_admin import auth as firebase_auth

        decoded_token = firebase_auth.verify_id_token(token)
        if decoded_token.get("role") != "SuperAdmin":
            raise HTTPException(status_code=403, detail="Solo SuperAdmins pueden traducir normas")

        org_ref = db.collection("organizations").document(request.organizationId)

        # 1. Collect English rules from both collections
        english_rules = []
        for col_name in ("rules", "pendingRules"):
            snap = org_ref.collection(col_name).stream()
            for doc_snap in snap:
                data = doc_snap.to_dict()
                name = data.get("name") or data.get("rule") or ""
                desc = data.get("description") or ""
                if _is_english_text(f"{name} {desc}"):
                    if (data.get("source") or "").startswith("RAE"):
                        continue  # never touch RAE rules
                    english_rules.append({
                        "ref": doc_snap.reference,
                        "name": name,
                        "description": desc,
                    })

        if not english_rules:
            return {"status": "ok", "translated": 0, "message": "No se encontraron normas en inglés."}

        logger.info(f"translate-rules: {len(english_rules)} English rules found in org={request.organizationId}")

        # 2. Translate in batches of 10 using the configured Gemini client
        BATCH_SIZE = 10
        schema = {
            "type": "ARRAY",
            "items": {
                "type": "OBJECT",
                "properties": {
                    "nombre": {"type": "STRING"},
                    "descripcion": {"type": "STRING"},
                },
                "required": ["nombre", "descripcion"],
            },
        }

        translated_count = 0
        for i in range(0, len(english_rules), BATCH_SIZE):
            batch_rules = english_rules[i:i + BATCH_SIZE]
            list_text = "\n".join(
                f'{j+1}. Nombre: "{r["name"]}" | Descripción: "{r["description"][:200]}"'
                for j, r in enumerate(batch_rules)
            )
            prompt = (
                "Eres un editor experto en lengua española. "
                "Traduce al español cada una de las siguientes normas editoriales.\n"
                "Para cada una devuelve:\n"
                "- \"nombre\": nombre breve de la norma en español (máx. 8 palabras)\n"
                "- \"descripcion\": explicación práctica en español (1-2 frases)\n\n"
                "IMPORTANTE: Toda la respuesta en español. No uses inglés en ningún campo.\n\n"
                f"{list_text}"
            )

            if client:
                try:
                    response = client.models.generate_content(
                        model=settings.LLM_MODEL,
                        contents=prompt,
                        config=types.GenerateContentConfig(
                            response_mime_type="application/json",
                            response_schema=schema,
                            temperature=0.2,
                        ),
                    )
                    translated = json.loads(response.text)
                except Exception as e:
                    logger.warning(f"Gemini translation error batch {i}: {e}")
                    translated = [{"nombre": r["name"], "descripcion": r["description"]} for r in batch_rules]
            else:
                translated = [{"nombre": r["name"], "descripcion": r["description"]} for r in batch_rules]

            # 3. Write translated names back to Firestore
            firestore_batch = db.batch()
            for j, rule in enumerate(batch_rules):
                t = translated[j] if j < len(translated) else {"nombre": rule["name"], "descripcion": rule["description"]}
                firestore_batch.update(rule["ref"], {
                    "name": t["nombre"],
                    "rule": t["nombre"],
                    "description": t["descripcion"],
                })
                translated_count += 1
            firestore_batch.commit()

        logger.info(f"translate-rules: translated {translated_count} rules for org={request.organizationId}")
        return {
            "status": "success",
            "organizationId": request.organizationId,
            "translated": translated_count,
        }

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error translating rules: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=str(e))


if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)
