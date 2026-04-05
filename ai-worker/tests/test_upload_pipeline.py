import pytest
import requests
import time
import os
from unittest.mock import patch, MagicMock

# The backend runs on localhost:8000
BASE_URL = "http://localhost:8000"

def test_ingest_book_endpoint_exists():
    """
    Test that the /api/v1/ingest-book endpoint exists 
    and returns a reasonable response (e.g. 422 if missing args, or 200).
    """
    # Send a request with no body
    try:
        response = requests.post(f"{BASE_URL}/api/v1/ingest-book", json={})
        # If it returns 422 Unprocessable Entity, it means validation failed, but endpoint exists
        # If it returns 500, something is wrong
        assert response.status_code in [422, 200], f"Unexpected status code: {response.status_code}"
    except requests.exceptions.ConnectionError:
        pytest.fail("Backend server is not running on localhost:8000")

@patch('app.services.document_parser.parse_docx')
def test_mocked_ingest_payload(mock_parse_docx, test_client):
    """
    Test the ingestion endpoint through TestClient (mocked Firebase / AI).
    We mock parse_docx to return dummy paragraphs without downloading from real Storage.
    """
    mock_parse_docx.return_value = ["Paragraph 1", "Paragraph 2", "Paragraph 3"]
    
    # We use test_client from conftest (which mocks firebase_admin and vertexai)
    payload = {
        "bookId": "test_book_123",
        "organizationId": "ZH3he4agLYMRqLwOWSLQ",
        "fileUrl": "gs://fake-bucket/fake-file.docx",
        "authorId": "test_author"
    }
    
    # Create a minimal valid docx in memory
    from docx import Document
    import io
    doc = Document()
    doc.add_paragraph("Paragraph 1")
    doc.add_paragraph("Paragraph 2")
    doc.add_paragraph("Paragraph 3")
    fake_docx_stream = io.BytesIO()
    doc.save(fake_docx_stream)
    
    # Patch requests.get so it doesn't try to download a real file
    with patch('requests.get') as mock_get:
        mock_resp = MagicMock()
        mock_resp.status_code = 200
        mock_resp.content = fake_docx_stream.getvalue()
        mock_get.return_value = mock_resp

        response = test_client.post("/api/v1/ingest-book", json=payload)
        
        # Check standard success payload
        assert response.status_code == 200
        assert response.json()["status"] == "success"
        assert response.json()["total_chunks"] == 3
